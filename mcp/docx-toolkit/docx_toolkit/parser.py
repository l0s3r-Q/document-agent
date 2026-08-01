"""parser.py —— 解析 .docx：页面设置、样式体系、结构树。"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
}

_LINE_RULE_MAP = {
    "EXACTLY": "EXACTLY",
    "AT_LEAST": "AT_LEAST",
    "MULTIPLE": "MULTIPLE",
}


def _run_font_info(run) -> dict:
    """提取 run 的中西文字体与字号。"""
    info = {"font_name": None, "east_asia": None, "size_pt": None, "bold": None}
    if run.font.name:
        info["font_name"] = run.font.name
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        east = rpr.rFonts.get(qn("w:eastAsia"))
        if east:
            info["east_asia"] = east
            if not info["font_name"]:
                info["font_name"] = east
    if run.font.size is not None:
        info["size_pt"] = run.font.size.pt
    info["bold"] = run.font.bold
    return info


def _paragraph_info(p) -> dict:
    """提取段落样式信息。"""
    pf = p.paragraph_format
    info = {
        "style_name": p.style.name if p.style else None,
        "align": _ALIGN_MAP.get(pf.alignment, "LEFT"),
        "line_spacing_rule": _LINE_RULE_MAP.get(str(pf.line_spacing_rule), "MULTIPLE"),
        "line_spacing_pt": None,
        "line_spacing_multiple": None,
        "first_line_indent_pt": None,
        "text": p.text,
    }
    if pf.line_spacing is not None:
        if info["line_spacing_rule"] == "MULTIPLE":
            info["line_spacing_multiple"] = pf.line_spacing
        else:
            try:
                info["line_spacing_pt"] = pf.line_spacing.pt
            except AttributeError:
                info["line_spacing_pt"] = None
    # 优先读 w:firstLineChars（字符数），否则读 w:firstLine 绝对值
    ind = p._p.pPr.ind if p._p.pPr is not None and p._p.pPr.ind is not None else None
    if ind is not None:
        chars = ind.get(qn("w:firstLineChars"))
        if chars:
            # firstLineChars 单位是 1/100 字符，转磅需结合字号
            info["first_line_indent_chars"] = int(chars) / 100.0
    if pf.first_line_indent is not None:
        try:
            info["first_line_indent_pt"] = pf.first_line_indent.pt
        except AttributeError:
            info["first_line_indent_pt"] = None
    # 取首个非空 run 的字体作为该段代表样式
    for run in p.runs:
        fi = _run_font_info(run)
        if fi["font_name"] or fi["size_pt"]:
            info.update(fi)
            break
    return info


def _detect_type(p) -> str:
    """判定段落类型：优先 w:outlineLvl，其次样式名。"""
    ppr = p._p.pPr
    if ppr is not None and ppr.outlineLvl is not None:
        lvl = ppr.outlineLvl.val
        if lvl in (0, 1, 2):
            return f"heading{lvl + 1}"
        if lvl >= 3:
            return "heading3"
    name = (p.style.name if p.style else "") or ""
    if name.startswith("Heading 1") or name == "标题 1":
        return "heading1"
    if name.startswith("Heading 2") or name == "标题 2":
        return "heading2"
    if name.startswith("Heading 3") or name == "标题 3":
        return "heading3"
    if name.startswith("Heading"):
        return "heading1"
    return "paragraph"


def parse(path: str) -> dict:
    """解析 docx 文件，返回页面/样式/结构信息。"""
    doc = Document(path)
    sec = doc.sections[0]
    page = {}
    for key, margin in (("top_cm", sec.top_margin), ("bottom_cm", sec.bottom_margin),
                        ("left_cm", sec.left_margin), ("right_cm", sec.right_margin)):
        if margin is not None:
            page[key] = round(margin.cm, 2)

    styles = []
    structure = []
    seen_styles = set()

    # 建立 XML 元素 → 段落对象映射（body 子元素顺序即文档排版顺序）
    para_by_el = {id(p._p): p for p in doc.paragraphs}
    table_iter = iter(doc.tables)

    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = para_by_el.get(id(child))
            if p is None:
                continue
            info = _paragraph_info(p)
            ptype = _detect_type(p)
            entry = {"type": ptype, "text": info.pop("text")}
            entry.update({k: v for k, v in info.items() if v is not None})
            structure.append(entry)
            sig = (ptype, info.get("font_name"), info.get("size_pt"))
            if sig not in seen_styles and info.get("font_name"):
                seen_styles.add(sig)
                styles.append(info)
        elif tag == "tbl":
            tbl = next(table_iter, None)
            if tbl is None:
                continue
            rows = [[cell.text for cell in row.cells] for row in tbl.rows]
            structure.append({"type": "table", "rows": rows, "cols": len(rows[0]) if rows else 0})
        elif tag in ("sdt", "fldSimple", "customXml"):
            # 内容控件/域/自定义 XML：结构树中标记，提示内容未展开
            structure.append({"type": "unsupported", "tag": tag})

    return {
        "file": path,
        "page": page,
        "styles": styles,
        "structure": structure,
    }


def extract_structure(path: str) -> dict:
    """仅提取结构树（标题+正文大纲），供"不改结构"场景使用。"""
    data = parse(path)
    tree = []
    for item in data["structure"]:
        if item["type"] in ("heading1", "heading2", "heading3"):
            tree.append({"level": int(item["type"][-1]), "text": item["text"]})
    return {
        "file": path,
        "outline": tree,
        "paragraph_count": len([s for s in data["structure"] if s["type"] == "paragraph"]),
    }
