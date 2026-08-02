"""builder.py —— 按 DocumentSpec 生成 .docx：字体/字号/行距/缩进/对齐/表格。"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .styles import BUILTIN_STYLES

_ALIGN = {
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_run_font(run, font_name: str | None, size_pt: float | None, bold: bool | None):
    """设置 run 字体（中西文都设置，中文必须写 w:eastAsia）。"""
    if font_name:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)
    if size_pt:
        run.font.size = Pt(float(size_pt))
    if bold is not None:
        run.font.bold = bold


def _apply_paragraph_format(pf, style: dict):
    """按样式字典设置段落格式。"""
    align = style.get("align")
    if align and align in _ALIGN:
        pf.alignment = _ALIGN[align]

    rule = style.get("line_spacing_rule")
    if rule == "EXACTLY" and style.get("line_spacing_pt") is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(float(style["line_spacing_pt"]))
    elif rule == "AT_LEAST" and style.get("line_spacing_pt") is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        pf.line_spacing = Pt(float(style["line_spacing_pt"]))
    elif style.get("line_spacing_multiple") is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = float(style["line_spacing_multiple"])

    if style.get("first_line_indent_pt") is not None:
        pf.first_line_indent = Pt(float(style["first_line_indent_pt"]))
    elif style.get("first_line_indent_chars") and style.get("size_pt"):
        # 缩进 N 字符 ≈ N × 字号
        pf.first_line_indent = Pt(float(style["size_pt"]) * float(style["first_line_indent_chars"]))


def _resolve_style(doc_type: str, role: str, override: dict | None = None, spec_styles: dict | None = None) -> dict:
    """取预置样式，叠加 spec.styles（按 role 覆盖）与 section 级覆盖。"""
    base = BUILTIN_STYLES.get(doc_type, BUILTIN_STYLES["general"])["roles"].get(role, {})
    merged = dict(base)
    if spec_styles and role in spec_styles:
        merged.update(spec_styles[role])
    if override:
        merged.update(override)
    return merged


def build(spec: dict, output_path: str) -> dict:
    """按 DocumentSpec 生成 docx。

    spec 契约:
      {doc_type, title, author?, date?, page{top_cm,bottom_cm,left_cm,right_cm},
       styles?{role: {font_name,size_pt,bold,align,line_spacing_rule,...}} 按 role 覆盖预置排版,
       sections:[{type,text,font?,align?,bold?,indent_first_line?,rows?,items?}]}
    """
    if not output_path.lower().endswith(".docx"):
        return {"ok": False, "error": "输出路径必须以 .docx 结尾"}
    doc = Document()
    doc_type = spec.get("doc_type", "general")
    warnings_list = []

    # ── 页面设置 ──────────────────────────────────────────────
    sec = doc.sections[0]
    page = spec.get("page") or BUILTIN_STYLES.get(doc_type, BUILTIN_STYLES["general"])["page"]
    sec.top_margin = Cm(page.get("top_cm", 2.54))
    sec.bottom_margin = Cm(page.get("bottom_cm", 2.54))
    sec.left_margin = Cm(page.get("left_cm", 3.17))
    sec.right_margin = Cm(page.get("right_cm", 3.17))

    # ── 页眉 / 页脚 / 页码（spec.header / spec.footer / spec.page_number）──
    if spec.get("header"):
        sec.header.is_linked_to_previous = False
        hp = sec.header.paragraphs[0]
        hp.text = ""
        hr = hp.add_run(spec["header"])
        _set_run_font(hr, "宋体", 10, False)
    if spec.get("footer"):
        sec.footer.is_linked_to_previous = False
        fp = sec.footer.paragraphs[0]
        fp.text = ""
        fr = fp.add_run(spec["footer"])
        _set_run_font(fr, "宋体", 10, False)
    if spec.get("page_number"):
        sec.footer.is_linked_to_previous = False
        pnum_p = sec.footer.paragraphs[0]
        pnum_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        pnum_p._p.append(fld)

    spec_styles = spec.get("styles") or {}

    def add_paragraph(sec_spec: dict, role: str):
        override = {}
        if sec_spec.get("font"):
            f = sec_spec["font"]
            if "name" in f and "font_name" not in f:
                f = dict(f, font_name=f["name"])
            if "east_asia" in f and "font_name" not in f:
                f = dict(f, font_name=f["east_asia"])
            override.update(f)
        if sec_spec.get("align"):
            override["align"] = sec_spec["align"]
        if sec_spec.get("bold") is not None:
            override["bold"] = sec_spec["bold"]
        style = _resolve_style(doc_type, role, override, spec_styles)

        p = doc.add_paragraph()
        # 标题设置 Word 内置 Heading 样式，确保大纲/结构可识别
        if role in ("heading1", "heading2", "heading3"):
            style_name = {"heading1": "Heading 1", "heading2": "Heading 2", "heading3": "Heading 3"}[role]
            try:
                p.style = doc.styles[style_name]
            except KeyError:
                pass
        _apply_paragraph_format(p.paragraph_format, style)
        text = sec_spec.get("text", "")
        if text:
            run = p.add_run(text)
            _set_run_font(
                run,
                style.get("font_name"),
                style.get("size_pt"),
                style.get("bold"),
            )
        return p

    # ── 标题 ──────────────────────────────────────────────────
    title = spec.get("title")
    if title:
        add_paragraph({"text": title}, "title")

    # ── 作者/日期行（可选）─────────────────────────────────────
    meta_line = "  ".join(x for x in (spec.get("author"), spec.get("date")) if x)
    if meta_line:
        add_paragraph({"text": meta_line, "align": "CENTER"}, "body")

    # ── 正文各节 ──────────────────────────────────────────────
    for item in spec.get("sections", []):
        stype = item.get("type", "paragraph")
        if stype == "toc":
            # 目录：TOC 域（Word 打开时按 F9 或自动更新；python-docx 无法直接渲染目录条目）
            p = doc.add_paragraph()
            fld_char_begin = OxmlElement("w:fldChar")
            fld_char_begin.set(qn("w:fldCharType"), "begin")
            instr_text = OxmlElement("w:instrText")
            instr_text.set(qn("xml:space"), "preserve")
            _bs = chr(92)  # 反斜杠（避免转义）
            instr_text.text = f" TOC {_bs}o \"1-3\" {_bs}h {_bs}z {_bs}u "
            fld_char_sep = OxmlElement("w:fldChar")
            fld_char_sep.set(qn("w:fldCharType"), "separate")
            t_el = OxmlElement("w:t")
            t_el.text = "目录将在打开文档时更新（Word 中按 F9 或点击更新域）"
            r_placeholder = OxmlElement("w:r")
            r_placeholder.append(t_el)
            fld_char_end = OxmlElement("w:fldChar")
            fld_char_end.set(qn("w:fldCharType"), "end")
            run = p.add_run()
            run._r.append(fld_char_begin)
            run2 = p.add_run()
            run2._r.append(instr_text)
            run3 = p.add_run()
            run3._r.append(fld_char_sep)
            p._p.append(r_placeholder)
            run4 = p.add_run()
            run4._r.append(fld_char_end)
        elif stype == "title":
            # 文档标题（居中大字段）
            add_paragraph({**item, "align": "CENTER"}, "title")
        elif stype in ("heading1", "heading2", "heading3"):
            add_paragraph(item, stype)
        elif stype == "paragraph":
            add_paragraph(item, "body")
        elif stype == "page_break":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
        elif stype == "separator":
            p = doc.add_paragraph()
            run = p.add_run("—" * 20)
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stype == "image":
            # 插入图片：{path, width_cm?, caption?}
            img_path = item.get("path", "")
            if img_path and os.path.exists(img_path):
                try:
                    from docx.shared import Cm as _Cm
                    width = _Cm(item.get("width_cm", 14))
                    doc.add_picture(img_path, width=width)
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:  # noqa: BLE001
                    warnings_list.append(f"图片插入失败: {type(e).__name__}")
            elif img_path:
                warnings_list.append(f"图片路径不存在: {img_path}")
            if item.get("caption"):
                add_paragraph({"text": item["caption"], "align": "CENTER"}, "body")

        elif stype == "list":
            style = _resolve_style(doc_type, "body", item.get("font"), spec_styles)
            for i, it in enumerate(item.get("items", []), 1):
                p = doc.add_paragraph()
                _apply_paragraph_format(p.paragraph_format, style)
                run = p.add_run(f"{i}. {it}")
                _set_run_font(run, style.get("font_name"), style.get("size_pt"), style.get("bold"))
        elif stype == "table":
            rows = item.get("rows", [])
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            style = _resolve_style(doc_type, "table", item.get("font"), spec_styles)
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    _apply_paragraph_format(p.paragraph_format, style)
                    run = p.add_run(row[ci] if ci < len(row) else "")
                    _set_run_font(run, style.get("font_name"), style.get("size_pt"), style.get("bold"))
            # 合并单元格：merges: [{from: "A1", to: "B2"}] 或 [{from: [r,c], to: [r,c]}]
            import re as _re
            def _coord(c):
                if isinstance(c, (list, tuple)) and len(c) == 2:
                    return (int(c[0]) - 1 if isinstance(c[0], int) else int(c[0]), int(c[1]) - 1 if isinstance(c[1], int) else int(c[1]))
                m = _re.fullmatch(r"([A-Za-z]+)([0-9]+)", str(c))
                if not m:
                    return None
                col = 0
                for ch in m.group(1).upper():
                    col = col * 26 + (ord(ch) - 64)
                return (int(m.group(2)) - 1, col - 1)
            for m in item.get("merges") or []:
                try:
                    c1, c2 = _coord(m["from"]), _coord(m["to"])
                    if c1 and c2:
                        table.cell(*c1).merge(table.cell(*c2))
                except (KeyError, IndexError, ValueError):
                    continue

    # 自动创建输出目录（不存在时）+ 原子写（临时文件 + replace，防中途损坏）
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    tmp_path = output_path + ".tmp"
    doc.save(tmp_path)
    os.replace(tmp_path, output_path)
    result = {"ok": True, "path": output_path}
    if warnings_list:
        result["warnings"] = warnings_list
    return result
